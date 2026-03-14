import fitz  # PyMuPDF for PDFs
import docx  # python-docx for Word
import re
import difflib
import os
import nltk
from nltk.tokenize import sent_tokenize

# Ensure NLTK is ready
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

class MultiDocSearch:
    def __init__(self):
        self.knowledge_base = []

    def add_documents(self, file_paths):
        """
        Accepts a list of file paths. Detects type and loads content.
        """
        for path in file_paths:
            if not os.path.exists(path):
                print(f"Skipping missing file: {path}")
                continue
            
            ext = path.lower().split('.')[-1]
            print(f"Processing {path}...")
            
            if ext == 'pdf':
                new_chunks = self._load_pdf(path)
            elif ext == 'docx':
                new_chunks = self._load_docx(path)
            else:
                print(f"Unsupported file type: {ext}")
                continue
                
            # Tag each chunk with its source filename
            for chunk in new_chunks:
                self.knowledge_base.append({
                    'text': chunk,
                    'source': os.path.basename(path)
                })
        
        print(f"Total knowledge base: {len(self.knowledge_base)} segments from {len(file_paths)} files.")

    def _clean_text(self, text):
        """Standardizes text (removes smart quotes, bullets, extra spaces)."""
        replacements = {
            "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
            "\u2013": "-", "\u2014": "-", "\u2026": "...", 
            "\uf0b7": "", "•": "", "●": "", "○": "",  
            "\u25cf": "", "\u25cb": "",        
            "\"": '"'                                
        }   
        for uni, std in replacements.items():
            text = text.replace(uni, std)
        text = text.replace('\n', ' ').strip()
        return re.sub(r'\s+', ' ', text)

    def _load_docx(self, path):
        doc = docx.Document(path)
        chunks = []
        # Paragraphs
        for para in doc.paragraphs:
            t = self._clean_text(para.text)
            if len(t) > 30: chunks.append(t)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = self._clean_text(cell.text)
                    if len(t) > 30: chunks.append(t)
        return chunks

    def _load_pdf(self, path):
        doc = fitz.open(path)
        chunks = []
        for page in doc:
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0])) # Sort top-down
            for b in blocks:
                t = self._clean_text(b[4])
                # Filter headers/footers
                if len(t) > 30 and "Page" not in t:
                    chunks.append(t)
        return chunks

    def _fuzzy_score(self, chunk, query):
        """Scoring logic (Exact Phrase > Fuzzy Word > Context Length)."""
        score = 0
        c_low = chunk.lower()
        q_low = query.lower()
        
        if q_low in c_low: score += 50
        
        q_words = q_low.split()
        c_words = c_low.split()
        
        for qw in q_words:
            if qw in c_low: 
                score += 10
            elif difflib.get_close_matches(qw, c_words, n=1, cutoff=0.85):
                score += 8
                
        return score

    def search(self, topics):
        """
        Searches ALL documents and aggregates the results for each topic.
        """
        final_report = {}

        for topic in topics:
            found_dates = set()
            snippets = []
            
            # Scan entire knowledge base
            for entry in self.knowledge_base:
                chunk = entry['text']
                source = entry['source']
                
                score = self._fuzzy_score(chunk, topic)
                
                # If we find a good match
                if score > 15:
                    # 1. Extract Dates
                    dates = re.findall(r'\b(19\d{2}|18\d{2})\b', chunk)
                    found_dates.update(dates)
                    
                    # 2. Extract Context (Definition sentence)
                    sentences = sent_tokenize(chunk)
                    best_sent = chunk[:200] # Fallback
                    for s in sentences:
                        # Prioritize sentences that look like definitions
                        if any(x in s.lower() for x in ["refers to", "was a", "passed", "led to", "known as"]):
                            best_sent = s
                            break
                    
                    # Store snippet with its source
                    snippets.append({
                        "file": source,
                        "text": best_sent
                    })

            # 3. Compile Results
            if snippets:
                # Deduplicate snippets based on text similarity (simple set check)
                unique_snippets = {s['text']: s for s in snippets}.values()
                
                final_report[topic] = {
                    "all_dates": sorted(list(found_dates)),
                    "sources_found": [s['file'] for s in unique_snippets],
                    "key_contexts": [s['text'] for s in unique_snippets]
                }
            else:
                final_report[topic] = {"error": "Not found in any document."}

        return final_report

# ==========================================
# EXAMPLE USAGE
# ==========================================
if __name__ == "__main__":
    # 1. Initialize
    app = MultiDocSearch()
    
    # 2. Add your files (Mix of PDF and Docx is fine)
    # Ensure these files exist in your folder
    my_files = [
        "src/notes_A.pdf", 
        "src/notes_B.pdf" 
    ]
    
    app.add_documents(my_files)

    # 3. Define Topics
    # topics = ["Enabling Act", "Nuremberg Laws"]
    topic_input = input("Enter a topic (eg. Enabling Act, Nuremberg Laws): ")
    topics = topic_input.split(',')

    # 4. Search & Print
    results = app.search(topics)
    
    import json
    print(json.dumps(results, indent=2))